import os
from typing import Dict, List
from nltk.tokenize import word_tokenize
import docx
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter


def docx_filename_to_url(filename: str) -> str:
    """
    Convert a filename like:
        help.unhcr.org_denmark_about-unhcr-in-denmark_.docx
    to a URL like:
        https://help.unhcr.org/denmark/about-unhcr-in-denmark/
    """
    name_no_ext = filename.replace(".docx", "")
    parts = [p for p in name_no_ext.split("_") if p]
    if not parts:
        return ""
    domain = parts[0]
    path_parts = parts[1:]
    url = "https://" + domain + "/"
    if path_parts:
        url += "/".join(path_parts) + "/"
    return url


def build_doc_url_dict(raw_folder: str) -> Dict[str, str]:
    """Build a mapping of .docx paths to their corresponding help site URLs."""
    doc_url_dict = {}
    for filename in os.listdir(raw_folder):
        if filename.endswith(".docx"):
            doc_path = os.path.join(raw_folder, filename)
            url = docx_filename_to_url(filename)
            doc_url_dict[doc_path] = url
    return doc_url_dict


def extract_text_from_docx(doc_path: str) -> str:
    """Extract and return cleaned text content from a .docx file."""
    document = docx.Document(doc_path)
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    return "\\n".join(paragraphs)


def load_docs_with_metadata(raw_folder: str) -> List[Document]:
    """Load DOCX files as LangChain Documents with metadata linking to source URLs."""
    doc_url_dict = build_doc_url_dict(raw_folder)
    docs = []
    for doc_path, url in doc_url_dict.items():
        text = extract_text_from_docx(doc_path)
        if text.strip():
            docs.append(Document(page_content=text, metadata={"source": url}))
    return docs


def chunk_documents(docs: List[Document], chunk_size=500, chunk_overlap=50):
    """Chunk parent documents into overlapping child chunks for retrieval."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    parent_docs, child_docs = [], []
    for doc in docs:
        parent_docs.append(doc)
        splitted_texts = text_splitter.split_text(doc.page_content)
        for chunk_text in splitted_texts:
            child_docs.append(Document(
                page_content=chunk_text,
                metadata={"parent_source": doc.metadata["source"]}
            ))
    return parent_docs, child_docs


def embed_child_chunks(child_docs: List[Document], embedding_model) -> List[List[float]]:
    """Generate vector embeddings for each child document using the provided embedding model."""
    texts = [child.page_content for child in child_docs]
    vectors = embedding_model.embed_documents(texts)
    return vectors

from pinecone import Pinecone
from sentence_transformers import CrossEncoder


def initialize_pinecone(api_key: str, index_name: str, dimension: int, environment="us-east-1-aws", region="us-east-1"):
    """Initialize Pinecone and create the index if it doesn't exist."""
    pc = Pinecone(api_key=api_key, environment=environment, region=region)
    if not pc.has_index(index_name):
        pc.create_index(index_name, dimension=dimension)
    return pc.Index(index_name)


def upsert_to_pinecone(index, child_docs: List[Document], vectors: List[List[float]], namespace="text_chunks"):
    """Upsert embedded child documents and their metadata into Pinecone."""
    to_upsert = []
    for i, child_doc in enumerate(child_docs):
        to_upsert.append((
            f"child_{i}",
            vectors[i],
            {
                "parent_source": child_doc.metadata.get("parent_source", ""),
                "text": child_doc.page_content
            }
        ))
    index.upsert(to_upsert, namespace=namespace)


def rerank_documents(query: str, retrieved_docs: List[dict], reranker_model: CrossEncoder):
    """Apply CrossEncoder reranking to a list of retrieved documents."""
    query_doc_pairs = [(query, doc["content"]) for doc in retrieved_docs]
    rerank_scores = reranker_model.predict(query_doc_pairs)
    for i, doc in enumerate(retrieved_docs):
        doc["rerank_score"] = rerank_scores[i]
    reranked_docs = sorted(retrieved_docs, key=lambda x: x["rerank_score"], reverse=True)
    return reranked_docs


def store_retrieve_and_rerank(query: str, index, embedding_model, reranker_model, top_k_retrieve=5, top_k_rerank=3):
    """Embed query, retrieve documents from Pinecone, rerank them, and return top results."""
    query_vector = embedding_model.embed_query(query)
    results = index.query(
        vector=query_vector,
        top_k=top_k_retrieve,
        includeMetadata=True,
        namespace="text_chunks"
    )
    retrieved_docs = []
    for match in results["matches"]:
        text_content = match["metadata"].get("text", "").strip()
        source_url = match["metadata"].get("parent_source", "").strip()
        if text_content:
            retrieved_docs.append({
                "id": match["id"],
                "score": match["score"],
                "content": text_content,
                "source": source_url
            })
    if not retrieved_docs:
        return []

    reranked_docs = rerank_documents(query, retrieved_docs, reranker_model)
    return reranked_docs[:top_k_rerank]

def generate_answer(query: str, reranked_docs: List[dict], prompt_template: str, llm_pipeline) -> str:
    """Generate a final answer using the top reranked docs and the Phi-3.5 pipeline."""
    context = "\\n\\n".join([doc["content"][:500] for doc in reranked_docs])
    sources = list({doc.get("source", "") for doc in reranked_docs if doc.get("source", "")})
    sources_str = ", ".join(sources) if sources else "N/A"

    formatted_prompt = prompt_template.format(context=context, query=query, sources=sources_str)
    response = llm_pipeline(formatted_prompt)

    if isinstance(response, list) and len(response) > 0:
        response_text = response[0]["generated_text"] if isinstance(response[0], dict) else response[0]
    else:
        response_text = response.strip() if isinstance(response, str) else str(response)

    return response_text.split("\\n\\n")[0].strip()



